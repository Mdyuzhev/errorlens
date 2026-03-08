"""Tests for article import service and endpoints (MD + DOCX)."""

import asyncio
import io
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from fastapi import HTTPException, UploadFile

from app.services.article_import_service import ArticleImportService


# ─── Fixtures ───────────────────────────────────────────────────────


def make_upload_file(
    content: bytes, filename: str, content_type: str = "application/octet-stream"
) -> UploadFile:
    """Create a mock UploadFile from bytes."""
    file_obj = io.BytesIO(content)
    return UploadFile(file=file_obj, filename=filename, headers=None)


@pytest.fixture
def sample_md_content() -> bytes:
    return b"# Test Article\n\nSome **bold** and *italic* content.\n\n## Section\n\n- item 1\n- item 2"


@pytest.fixture
def sample_md_no_heading() -> bytes:
    return b"This is a plain text article without a heading.\n\nSecond paragraph."


@pytest.fixture
def sample_cyrillic_md() -> bytes:
    return "# Тестовая статья\n\nСодержимое на русском.".encode("utf-8")


@pytest.fixture
def sample_bom_md() -> bytes:
    return "\ufeff# BOM Article\n\nContent after BOM.".encode("utf-8")


@pytest.fixture
def sample_docx_bytes(tmp_path: Path) -> bytes:
    """Generate a sample DOCX file with headings, bold, lists."""
    from docx import Document

    doc = Document()
    doc.add_heading("Test DOCX Article", level=1)
    p = doc.add_paragraph()
    run = p.add_run("This is bold text.")
    run.bold = True
    p2 = doc.add_paragraph()
    run2 = p2.add_run("This is italic text.")
    run2.italic = True
    doc.add_heading("Sub Section", level=2)
    doc.add_paragraph("Item one", style="List Bullet")
    doc.add_paragraph("Item two", style="List Bullet")
    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return path.read_bytes()


@pytest.fixture
def sample_docx_with_table(tmp_path: Path) -> bytes:
    """Generate a DOCX with a table."""
    from docx import Document

    doc = Document()
    doc.add_heading("Table Article", level=1)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"
    path = tmp_path / "table.docx"
    doc.save(str(path))
    return path.read_bytes()


@pytest.fixture
def empty_file_bytes() -> bytes:
    return b""


@pytest.fixture
def mock_db():
    """Mock AsyncSession."""
    db = AsyncMock()
    return db


# ─── Service Unit Tests ────────────────────────────────────────────


class TestArticleImportService:
    """Unit tests for ArticleImportService parsing logic."""

    def test_parse_markdown_basic(self, sample_md_content: bytes):
        service = ArticleImportService.__new__(ArticleImportService)
        title, content = service._parse_markdown(sample_md_content)
        assert title == "Test Article"
        assert "**bold**" in content
        assert "*italic*" in content

    def test_parse_markdown_extracts_title(self, sample_md_content: bytes):
        service = ArticleImportService.__new__(ArticleImportService)
        title, _ = service._parse_markdown(sample_md_content)
        assert title == "Test Article"

    def test_parse_markdown_no_heading(self, sample_md_no_heading: bytes):
        service = ArticleImportService.__new__(ArticleImportService)
        title, content = service._parse_markdown(sample_md_no_heading)
        assert title == "This is a plain text article without a heading."
        assert "Second paragraph." in content

    def test_parse_markdown_utf8_cyrillic(self, sample_cyrillic_md: bytes):
        service = ArticleImportService.__new__(ArticleImportService)
        title, content = service._parse_markdown(sample_cyrillic_md)
        assert title == "Тестовая статья"
        assert "Содержимое на русском." in content

    def test_parse_markdown_bom(self, sample_bom_md: bytes):
        service = ArticleImportService.__new__(ArticleImportService)
        title, content = service._parse_markdown(sample_bom_md)
        assert title == "BOM Article"
        assert not content.startswith("\ufeff")

    def test_parse_markdown_latin1_fallback(self):
        # latin-1 encoded content (not valid utf-8)
        content = "# Résumé\n\nContent with accents.".encode("latin-1")
        service = ArticleImportService.__new__(ArticleImportService)
        title, text = service._parse_markdown(content)
        assert "sum" in title.lower()

    def test_parse_docx_basic(self, sample_docx_bytes: bytes):
        service = ArticleImportService.__new__(ArticleImportService)
        warnings: list[str] = []
        title, markdown = service._parse_docx(sample_docx_bytes, warnings)
        assert title == "Test DOCX Article"
        assert len(markdown) > 0

    def test_parse_docx_preserves_headings(self, sample_docx_bytes: bytes):
        service = ArticleImportService.__new__(ArticleImportService)
        warnings: list[str] = []
        _, markdown = service._parse_docx(sample_docx_bytes, warnings)
        assert "Sub Section" in markdown

    def test_parse_docx_preserves_bold_italic(self, sample_docx_bytes: bytes):
        service = ArticleImportService.__new__(ArticleImportService)
        warnings: list[str] = []
        _, markdown = service._parse_docx(sample_docx_bytes, warnings)
        assert "bold" in markdown.lower()
        assert "italic" in markdown.lower()

    def test_parse_docx_preserves_lists(self, sample_docx_bytes: bytes):
        service = ArticleImportService.__new__(ArticleImportService)
        warnings: list[str] = []
        _, markdown = service._parse_docx(sample_docx_bytes, warnings)
        assert "Item one" in markdown
        assert "Item two" in markdown

    def test_parse_docx_tables(self, sample_docx_with_table: bytes):
        service = ArticleImportService.__new__(ArticleImportService)
        warnings: list[str] = []
        title, markdown = service._parse_docx(sample_docx_with_table, warnings)
        assert title == "Table Article"
        assert "A1" in markdown
        assert "B1" in markdown

    def test_extract_title_heading(self):
        service = ArticleImportService.__new__(ArticleImportService)
        assert service._extract_title("# My Title\n\nContent") == "My Title"

    def test_extract_title_no_heading(self):
        service = ArticleImportService.__new__(ArticleImportService)
        assert service._extract_title("Just text\n\nMore text") == "Just text"

    def test_extract_title_empty(self):
        service = ArticleImportService.__new__(ArticleImportService)
        assert service._extract_title("") == "Untitled"
        assert service._extract_title("   \n  \n  ") == "Untitled"

    def test_sanitize_content_removes_scripts(self):
        service = ArticleImportService.__new__(ArticleImportService)
        result = service._sanitize_content("Hello <script>alert('xss')</script> World")
        assert "<script>" not in result
        assert "Hello" in result
        assert "World" in result

    def test_sanitize_content_removes_iframes(self):
        service = ArticleImportService.__new__(ArticleImportService)
        result = service._sanitize_content('Text <iframe src="evil"></iframe> more')
        assert "<iframe" not in result

    def test_sanitize_content_normalizes_newlines(self):
        service = ArticleImportService.__new__(ArticleImportService)
        result = service._sanitize_content("a\r\nb\rc\n\n\n\nd")
        assert "\r" not in result
        assert "\n\n\n" not in result

    def test_empty_input(self):
        """test_empty_input: empty markdown returns Untitled."""
        service = ArticleImportService.__new__(ArticleImportService)
        title, content = service._parse_markdown(b"")
        assert title == "Untitled"
        assert content == ""

    def test_none_handling(self):
        """test_none_handling: extract_title handles edge cases."""
        service = ArticleImportService.__new__(ArticleImportService)
        assert service._extract_title("") == "Untitled"

    def test_duplicate_handling(self):
        """test_duplicate_handling: multiple headings — uses first."""
        service = ArticleImportService.__new__(ArticleImportService)
        text = "# First\n\n# Second\n"
        assert service._extract_title(text) == "First"


# ─── Async Validation Tests ────────────────────────────────────────


class TestFileValidation:
    """Test file validation in _read_and_validate."""

    @pytest.mark.asyncio
    async def test_wrong_extension(self):
        upload = make_upload_file(b"some content", "file.txt")
        service = ArticleImportService.__new__(ArticleImportService)
        with pytest.raises(HTTPException) as exc_info:
            await service._read_and_validate(upload)
        assert exc_info.value.status_code == 400
        assert "Unsupported file format" in exc_info.value.detail

    @pytest.mark.asyncio
    async def test_empty_file(self):
        upload = make_upload_file(b"", "empty.md")
        service = ArticleImportService.__new__(ArticleImportService)
        with pytest.raises(HTTPException) as exc_info:
            await service._read_and_validate(upload)
        assert exc_info.value.status_code == 400
        assert "empty" in exc_info.value.detail.lower()

    @pytest.mark.asyncio
    async def test_too_large_file(self):
        big_content = b"x" * (5 * 1024 * 1024 + 1)
        upload = make_upload_file(big_content, "big.md")
        service = ArticleImportService.__new__(ArticleImportService)
        with pytest.raises(HTTPException) as exc_info:
            await service._read_and_validate(upload)
        assert exc_info.value.status_code == 413

    @pytest.mark.asyncio
    async def test_valid_md_file(self, sample_md_content: bytes):
        upload = make_upload_file(sample_md_content, "article.md")
        service = ArticleImportService.__new__(ArticleImportService)
        content, ext, warnings = await service._read_and_validate(upload)
        assert ext == ".md"
        assert content == sample_md_content
        assert isinstance(warnings, list)

    @pytest.mark.asyncio
    async def test_valid_docx_file(self, sample_docx_bytes: bytes):
        upload = make_upload_file(sample_docx_bytes, "article.docx")
        service = ArticleImportService.__new__(ArticleImportService)
        content, ext, warnings = await service._read_and_validate(upload)
        assert ext == ".docx"
        assert content == sample_docx_bytes

    @pytest.mark.asyncio
    async def test_pdf_extension_rejected(self):
        upload = make_upload_file(b"fake pdf content", "file.pdf")
        service = ArticleImportService.__new__(ArticleImportService)
        with pytest.raises(HTTPException) as exc_info:
            await service._read_and_validate(upload)
        assert exc_info.value.status_code == 400

    @pytest.mark.asyncio
    async def test_no_extension_rejected(self):
        upload = make_upload_file(b"content", "noextension")
        service = ArticleImportService.__new__(ArticleImportService)
        with pytest.raises(HTTPException) as exc_info:
            await service._read_and_validate(upload)
        assert exc_info.value.status_code == 400


# ─── Preview Tests ──────────────────────────────────────────────────


class TestPreviewFile:
    """Test preview_file returns parsed content without creating article."""

    @pytest.mark.asyncio
    async def test_preview_md(self, sample_md_content: bytes):
        upload = make_upload_file(sample_md_content, "article.md")
        service = ArticleImportService(None)  # type: ignore[arg-type]
        title, content, warnings = await service.preview_file(upload)
        assert title == "Test Article"
        assert "bold" in content
        assert isinstance(warnings, list)

    @pytest.mark.asyncio
    async def test_preview_docx(self, sample_docx_bytes: bytes):
        upload = make_upload_file(sample_docx_bytes, "article.docx")
        service = ArticleImportService(None)  # type: ignore[arg-type]
        title, content, warnings = await service.preview_file(upload)
        assert title == "Test DOCX Article"
        assert len(content) > 0

    @pytest.mark.asyncio
    async def test_preview_cyrillic(self, sample_cyrillic_md: bytes):
        upload = make_upload_file(sample_cyrillic_md, "ru.md")
        service = ArticleImportService(None)  # type: ignore[arg-type]
        title, content, warnings = await service.preview_file(upload)
        assert title == "Тестовая статья"
        assert "русском" in content


# ─── Import (Full Pipeline) Tests ──────────────────────────────────


class TestImportFromFile:
    """Test import_from_file with mocked ArticleService."""

    @pytest.mark.asyncio
    async def test_import_md_file(self, sample_md_content: bytes, mock_db):
        upload = make_upload_file(sample_md_content, "test.md")

        mock_article = MagicMock()
        mock_article.id = "art-1"
        mock_article.title = "Test Article"
        mock_article.slug = "test-article"
        mock_article.content = sample_md_content.decode()

        with patch.object(
            ArticleImportService, "__init__", lambda self, db: None
        ):
            service = ArticleImportService(mock_db)
            service.db = mock_db
            service.article_service = MagicMock()
            service.article_service.create_article = AsyncMock(return_value=mock_article)

            article, warnings = await service.import_from_file(
                file=upload,
                folder_id=None,
                project_id="proj-1",
                author="testuser",
                created_by="user-1",
            )

            assert article.title == "Test Article"
            service.article_service.create_article.assert_called_once()
            call_kwargs = service.article_service.create_article.call_args[1]
            assert call_kwargs["project_id"] == "proj-1"
            assert call_kwargs["author"] == "testuser"

    @pytest.mark.asyncio
    async def test_import_with_folder_id(self, sample_md_content: bytes, mock_db):
        upload = make_upload_file(sample_md_content, "test.md")

        mock_article = MagicMock()
        mock_article.id = "art-2"

        with patch.object(
            ArticleImportService, "__init__", lambda self, db: None
        ):
            service = ArticleImportService(mock_db)
            service.db = mock_db
            service.article_service = MagicMock()
            service.article_service.create_article = AsyncMock(return_value=mock_article)

            await service.import_from_file(
                file=upload,
                folder_id="folder-1",
                project_id="proj-1",
                author="testuser",
                created_by="user-1",
            )

            call_kwargs = service.article_service.create_article.call_args[1]
            assert call_kwargs["folder_id"] == "folder-1"

    @pytest.mark.asyncio
    async def test_import_with_category_tags(self, sample_md_content: bytes, mock_db):
        upload = make_upload_file(sample_md_content, "test.md")

        mock_article = MagicMock()
        mock_article.id = "art-3"

        with patch.object(
            ArticleImportService, "__init__", lambda self, db: None
        ):
            service = ArticleImportService(mock_db)
            service.db = mock_db
            service.article_service = MagicMock()
            service.article_service.create_article = AsyncMock(return_value=mock_article)

            await service.import_from_file(
                file=upload,
                folder_id=None,
                project_id="proj-1",
                author="testuser",
                created_by="user-1",
                category="API",
                tags=["testing", "guide"],
            )

            call_kwargs = service.article_service.create_article.call_args[1]
            assert call_kwargs["category"] == "API"
            assert call_kwargs["tags"] == ["testing", "guide"]


# ─── Concurrent Access Test ─────────────────────────────────────────


class TestConcurrentAccess:
    """test_concurrent_access: multiple parallel previews."""

    @pytest.mark.asyncio
    async def test_concurrent_preview(self, sample_md_content: bytes):
        async def do_preview():
            upload = make_upload_file(sample_md_content, "test.md")
            service = ArticleImportService(None)  # type: ignore[arg-type]
            return await service.preview_file(upload)

        results = await asyncio.gather(*[do_preview() for _ in range(5)])
        assert len(results) == 5
        for title, content, warnings in results:
            assert title == "Test Article"


# ─── Memory Cleanup Test ────────────────────────────────────────────


class TestMemoryCleanup:
    """test_memory_cleanup: service doesn't retain state between calls."""

    @pytest.mark.asyncio
    async def test_memory_cleanup(self, sample_md_content: bytes, sample_cyrillic_md: bytes):
        service = ArticleImportService(None)  # type: ignore[arg-type]

        upload1 = make_upload_file(sample_md_content, "first.md")
        title1, content1, _ = await service.preview_file(upload1)

        upload2 = make_upload_file(sample_cyrillic_md, "second.md")
        title2, content2, _ = await service.preview_file(upload2)

        assert title1 != title2
        assert content1 != content2


# ─── Error Recovery Test ────────────────────────────────────────────


class TestErrorRecovery:
    """test_error_recovery: service recovers after errors."""

    @pytest.mark.asyncio
    async def test_error_recovery(self, sample_md_content: bytes):
        service = ArticleImportService(None)  # type: ignore[arg-type]

        # First call: invalid file → error
        bad_upload = make_upload_file(b"content", "bad.txt")
        with pytest.raises(HTTPException):
            await service.preview_file(bad_upload)

        # Second call: valid file → success
        good_upload = make_upload_file(sample_md_content, "good.md")
        title, content, _ = await service.preview_file(good_upload)
        assert title == "Test Article"


# ─── Integration Tests (API endpoints) ─────────────────────────────


class TestArticleImportEndpoints:
    """Integration tests for /articles/import and /articles/import/preview."""

    def _get_auth_headers(self, client) -> dict:
        """Login as owner1 and return auth headers."""
        resp = client.post(
            "/auth/login", json={"username": "owner1", "password": "Test123!"}
        )
        if resp.status_code != 200:
            pytest.skip("Cannot login as owner1")
        return {"Authorization": f"Bearer {resp.json()['access_token']}"}

    def test_import_md_endpoint(self, client):
        headers = self._get_auth_headers(client)
        content = b"# Import MD Endpoint Test\n\nSome **bold** content."
        resp = client.post(
            "/articles/import",
            files={"file": ("test.md", io.BytesIO(content), "text/markdown")},
            data={"status": "draft"},
            headers=headers,
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["title"] == "Import MD Endpoint Test"
        assert "id" in data
        assert "slug" in data
        assert data["content_length"] > 0

    def test_import_docx_endpoint(self, client, sample_docx_bytes: bytes):
        headers = self._get_auth_headers(client)
        resp = client.post(
            "/articles/import",
            files={"file": ("test.docx", io.BytesIO(sample_docx_bytes), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"status": "draft"},
            headers=headers,
        )
        assert resp.status_code == 200
        data = resp.json()
        assert "Test DOCX Article" in data["title"]

    def test_import_wrong_extension(self, client):
        headers = self._get_auth_headers(client)
        resp = client.post(
            "/articles/import",
            files={"file": ("file.txt", io.BytesIO(b"text"), "text/plain")},
            headers=headers,
        )
        assert resp.status_code == 400
        assert "Unsupported" in resp.json()["detail"]

    def test_import_empty_file(self, client):
        headers = self._get_auth_headers(client)
        resp = client.post(
            "/articles/import",
            files={"file": ("empty.md", io.BytesIO(b""), "text/markdown")},
            headers=headers,
        )
        assert resp.status_code == 400
        assert "empty" in resp.json()["detail"].lower()

    def test_import_too_large(self, client):
        headers = self._get_auth_headers(client)
        big = b"x" * (5 * 1024 * 1024 + 1)
        resp = client.post(
            "/articles/import",
            files={"file": ("big.md", io.BytesIO(big), "text/markdown")},
            headers=headers,
        )
        assert resp.status_code == 413

    def test_import_unauthorized(self, client):
        resp = client.post(
            "/articles/import",
            files={"file": ("test.md", io.BytesIO(b"# Title"), "text/markdown")},
        )
        assert resp.status_code in (401, 403)

    def test_preview_md_endpoint(self, client, sample_md_content: bytes):
        headers = self._get_auth_headers(client)
        resp = client.post(
            "/articles/import/preview",
            files={"file": ("test.md", io.BytesIO(sample_md_content), "text/markdown")},
            headers=headers,
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["title"] == "Test Article"
        assert "bold" in data["content"]
        assert isinstance(data["warnings"], list)

    def test_preview_docx_endpoint(self, client, sample_docx_bytes: bytes):
        headers = self._get_auth_headers(client)
        resp = client.post(
            "/articles/import/preview",
            files={"file": ("test.docx", io.BytesIO(sample_docx_bytes), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            headers=headers,
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["title"] == "Test DOCX Article"
        assert len(data["content"]) > 0

    def test_import_with_nonexistent_folder_id(self, client):
        headers = self._get_auth_headers(client)
        content = b"# Folder Test Article\n\nContent for folder test."
        resp = client.post(
            "/articles/import",
            files={"file": ("folder.md", io.BytesIO(content), "text/markdown")},
            data={"folder_id": "nonexistent-folder", "status": "draft"},
            headers=headers,
        )
        # FK constraint prevents insertion with non-existent folder_id
        assert resp.status_code in (400, 500)

    def test_import_with_category_tags(self, client):
        headers = self._get_auth_headers(client)
        content = b"# Category Tags Article\n\nContent for category test."
        resp = client.post(
            "/articles/import",
            files={"file": ("cat.md", io.BytesIO(content), "text/markdown")},
            data={"category": "QA", "tags": "api,testing", "status": "draft"},
            headers=headers,
        )
        assert resp.status_code == 200

    def test_import_cyrillic_md(self, client):
        headers = self._get_auth_headers(client)
        content = "# Импорт кириллической статьи\n\nСодержимое на русском.".encode("utf-8")
        resp = client.post(
            "/articles/import",
            files={"file": ("ru.md", io.BytesIO(content), "text/markdown")},
            headers=headers,
        )
        assert resp.status_code == 200
        assert resp.json()["title"] == "Импорт кириллической статьи"
